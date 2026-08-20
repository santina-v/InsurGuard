import io
import csv

import fitz
from PIL import Image
import pytesseract
from docx import Document
from openpyxl import load_workbook


# ============================================================
# SUPPORTED FILE TYPES
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".jpg",
    ".jpeg",
    ".png",
    ".txt",
    ".xlsx",
    ".xls",
    ".csv",
}


# ============================================================
# PDF
# ============================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a text-based PDF.
    """

    try:
        document = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        text_parts = []

        for page in document:
            text = page.get_text()

            if text:
                text_parts.append(text.strip())

        document.close()

        extracted_text = "\n\n".join(
            text_parts
        ).strip()

        if not extracted_text:
            raise ValueError(
                "No readable text found in the PDF."
            )

        return extracted_text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from PDF: {exc}"
        ) from exc


# ============================================================
# DOCX
# ============================================================

def extract_text_from_docx(
    file_bytes: bytes
) -> str:

    try:
        document = Document(
            io.BytesIO(file_bytes)
        )

        paragraphs = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        # Also extract table contents
        for table in document.tables:
            for row in table.rows:

                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    paragraphs.append(
                        " | ".join(row_text)
                    )

        extracted_text = "\n".join(
            paragraphs
        ).strip()

        if not extracted_text:
            raise ValueError(
                "No readable text found in DOCX."
            )

        return extracted_text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from DOCX: {exc}"
        ) from exc


# ============================================================
# DOC
# ============================================================

def extract_text_from_doc(
    file_bytes: bytes
) -> str:

    """
    Legacy .doc files are harder to process directly.

    This implementation attempts to decode readable
    text from the legacy Word binary format.
    """

    try:
        # Try common encodings first
        for encoding in [
            "utf-8",
            "utf-16",
            "latin-1",
        ]:

            try:
                text = file_bytes.decode(
                    encoding,
                    errors="ignore"
                )

                # Remove excessive binary characters
                cleaned = "".join(
                    char
                    if char.isprintable() or char in "\n\r\t"
                    else " "
                    for char in text
                )

                lines = [
                    line.strip()
                    for line in cleaned.splitlines()
                    if line.strip()
                ]

                result = "\n".join(lines)

                if len(result) > 20:
                    return result

            except Exception:
                continue

        raise ValueError(
            "Unable to extract readable text from DOC file. "
            "Please convert it to DOCX or PDF."
        )

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from DOC: {exc}"
        ) from exc


# ============================================================
# IMAGE OCR
# ============================================================

def extract_text_from_image(
    file_bytes: bytes
) -> str:

    try:

        image = Image.open(
            io.BytesIO(file_bytes)
        )

        # Convert to RGB for OCR
        image = image.convert("RGB")

        text = pytesseract.image_to_string(
            image
        )

        extracted_text = text.strip()

        if not extracted_text:
            raise ValueError(
                "No readable text found in the image."
            )

        return extracted_text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from image: {exc}"
        ) from exc


# ============================================================
# TXT
# ============================================================

def extract_text_from_txt(
    file_bytes: bytes
) -> str:

    try:

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        ).strip()

        if not text:
            raise ValueError(
                "The TXT file is empty."
            )

        return text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to read TXT file: {exc}"
        ) from exc


# ============================================================
# EXCEL
# ============================================================

def extract_text_from_excel(
    file_bytes: bytes,
    extension: str
) -> str:

    try:

        workbook = load_workbook(
            filename=io.BytesIO(file_bytes),
            data_only=True
        )

        text_parts = []

        for worksheet in workbook.worksheets:

            text_parts.append(
                f"Sheet: {worksheet.title}"
            )

            for row in worksheet.iter_rows(
                values_only=True
            ):

                values = []

                for value in row:

                    if value is not None:
                        values.append(
                            str(value).strip()
                        )

                if values:
                    text_parts.append(
                        " | ".join(values)
                    )

        extracted_text = "\n".join(
            text_parts
        ).strip()

        if not extracted_text:
            raise ValueError(
                "No readable data found in Excel file."
            )

        return extracted_text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from Excel file: {exc}"
        ) from exc


# ============================================================
# CSV
# ============================================================

def extract_text_from_csv(
    file_bytes: bytes
) -> str:

    try:

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        )

        reader = csv.reader(
            io.StringIO(text)
        )

        rows = []

        for row in reader:

            values = [
                value.strip()
                for value in row
                if value.strip()
            ]

            if values:
                rows.append(
                    " | ".join(values)
                )

        extracted_text = "\n".join(
            rows
        ).strip()

        if not extracted_text:
            raise ValueError(
                "No readable data found in CSV file."
            )

        return extracted_text

    except ValueError:
        raise

    except Exception as exc:
        raise ValueError(
            f"Unable to extract text from CSV: {exc}"
        ) from exc


# ============================================================
# UNIVERSAL DOCUMENT EXTRACTOR
# ============================================================

def extract_text_from_document(
    file_bytes: bytes,
    filename: str
) -> str:

    if not filename:
        raise ValueError(
            "Uploaded file has no filename."
        )

    filename_lower = filename.lower()

    # Get extension
    extension = ""

    if "." in filename_lower:
        extension = "." + filename_lower.split(".")[-1]

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    if extension == ".pdf":
        return extract_text_from_pdf(
            file_bytes
        )

    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    if extension == ".docx":
        return extract_text_from_docx(
            file_bytes
        )

    # --------------------------------------------------------
    # DOC
    # --------------------------------------------------------

    if extension == ".doc":
        return extract_text_from_doc(
            file_bytes
        )

    # --------------------------------------------------------
    # IMAGES
    # --------------------------------------------------------

    if extension in {
        ".jpg",
        ".jpeg",
        ".png",
    }:

        return extract_text_from_image(
            file_bytes
        )

    # --------------------------------------------------------
    # TXT
    # --------------------------------------------------------

    if extension == ".txt":
        return extract_text_from_txt(
            file_bytes
        )

    # --------------------------------------------------------
    # EXCEL
    # --------------------------------------------------------

    if extension in {
        ".xlsx",
        ".xls",
    }:

        return extract_text_from_excel(
            file_bytes,
            extension
        )

    # --------------------------------------------------------
    # CSV
    # --------------------------------------------------------

    if extension == ".csv":
        return extract_text_from_csv(
            file_bytes
        )

    raise ValueError(
        f"Unsupported file format: {extension}"
    )